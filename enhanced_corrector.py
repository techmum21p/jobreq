"""
Enhanced Document Corrector
Generates properly formatted DOCX files that match the Albertsons requisition template

This agent uses the docx skill to create professional, correctly formatted documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List
import anthropic


class EnhancedDocumentCorrector:
    """
    Enhanced corrector that generates properly formatted Word documents
    Uses the docx library to maintain professional formatting
    """
    
    def __init__(self, api_key: str, template_path: str = None):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.template_path = template_path
    
    def generate_corrected_document(
        self,
        job_data: Dict,
        ground_truth: 'GroundTruth',
        pay_transparency_text: str,
        output_path: str
    ) -> str:
        """
        Generate a properly formatted requisition document with corrections
        
        Args:
            job_data: Extracted job data
            ground_truth: Ground truth corrections
            pay_transparency_text: Correct pay transparency text for the state
            output_path: Where to save the corrected document
            
        Returns:
            Path to the generated document
        """
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add Albertsons header
        header = doc.add_heading('Albertsons Companies', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Job Title
        title = doc.add_heading(job_data['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Location
        location = doc.add_paragraph(f"{job_data['location']}")
        location.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # JOB INFO Section
        doc.add_heading('JOB INFO', level=2)
        
        info_table_data = [
            ['Job Identification', job_data['job_id']],
            ['Job Category', job_data['category']],
            ['Posting Date', job_data['posting_date']],
            ['Job Schedule', job_data['schedule']],
            ['Locations', job_data['location']],
            ['Banner', job_data['banner']],
            ['Minimum Pay Rate', f"${ground_truth.min_pay_rate:.2f}"],  # CORRECTED
            ['Maximum Pay Rate', f"${ground_truth.max_pay_rate:.2f}"],  # CORRECTED
        ]
        
        for label, value in info_table_data:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        
        doc.add_paragraph()
        
        # JOB DESCRIPTION Section
        doc.add_heading('JOB DESCRIPTION', level=2)
        
        doc.add_heading('A Day in the Life:', level=3)
        doc.add_paragraph(job_data.get('day_in_life', ''))
        
        doc.add_heading('What you bring to the table:', level=3)
        for item in job_data.get('qualifications', []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading('Why you will choose us:', level=3)
        for item in job_data.get('benefits', []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Role-specific pay rates (if applicable)
        if job_data.get('role_rates'):
            doc.add_paragraph()
            doc.add_heading('Starting hourly rates:', level=3)
            
            doc.add_paragraph(
                "Starting hourly rates will be no less than the local minimum wage "
                "and will vary based on things like location, experience, qualifications, "
                "and the terms of any applicable collective bargaining agreement."
            )
            
            for role, (min_rate, max_rate) in job_data['role_rates'].items():
                # Ensure rates fall within ground truth range
                adjusted_min = max(min_rate, ground_truth.min_pay_rate)
                adjusted_max = min(max_rate, ground_truth.max_pay_rate)
                
                doc.add_paragraph(
                    f"{role} - ${adjusted_min:.2f} – ${adjusted_max:.2f}",
                    style='List Bullet'
                )
        
        doc.add_page_break()
        
        # ABOUT US Section
        doc.add_heading('ABOUT US', level=2)
        doc.add_paragraph(
            "Albertsons Companies is at the forefront of the revolution in retail. "
            "Committed to innovation and fostering a culture of belonging, our team is "
            "united with a unique purpose: to bring people together around the joys of "
            "food and to inspire well-being."
        )
        
        # Company details
        doc.add_paragraph(
            "Locally great and nationally strong, Albertsons Companies (NYSE: ACI) is "
            "a leading food and drug retailer in the U.S. We operate over 2,200 stores, "
            "1,732 pharmacies, 405 fuel centers, 22 distribution facilities, and 19 "
            "manufacturing plants across 34 states and the District of Columbia."
        )
        
        doc.add_paragraph()
        
        # Pay Transparency Section - CORRECTED
        doc.add_heading('Pay Transparency:', level=3)
        doc.add_paragraph(pay_transparency_text)
        
        # Disclaimer
        doc.add_paragraph()
        doc.add_heading('Disclaimer', level=3)
        doc.add_paragraph(
            "The above statements are intended to describe the general nature of work "
            "performed by the employees assigned to this job and are not the official "
            "job description for the position."
        )
        
        # Equal Opportunity Statement
        doc.add_paragraph()
        doc.add_heading('Albertsons is an Equal Opportunity Employer', level=3)
        doc.add_paragraph(
            "This Company is an Equal Opportunity Employer, and does not discriminate "
            "on the basis of race, gender, ethnicity, religion, national origin, age, "
            "disability, veteran status, gender identity/expression, sexual orientation, "
            "or on any other basis prohibited by law."
        )
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def extract_job_data_with_llm(self, pdf_text: str) -> Dict:
        """
        Use Claude to extract and structure job data from PDF text
        """
        
        prompt = """
        Extract the following information from this job requisition to populate a template:
        
        Required fields:
        - title: Job title
        - job_id: Job identification number
        - category: Job category
        - posting_date: Posting date
        - schedule: Job schedule (Full/Part time)
        - location: Complete address
        - banner: Company banner (e.g., Jewel Osco)
        - day_in_life: "A Day in the Life" description (2-3 paragraphs)
        - qualifications: List of "What you bring to the table" bullet points
        - benefits: List of "Why you will choose us" bullet points
        - role_rates: Dictionary of role names to [min, max] pay rates (if multiple roles listed)
        
        Return as JSON.
        
        Document:
        {text}
        """
        
        response = self.client.messages.create(
            model=self.model,
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": prompt.format(text=pdf_text)
            }]
        )
        
        return eval(response.content[0].text)  # Parse JSON response


class SmartCorrectionAgent:
    """
    Intelligent correction agent that determines which corrections to apply
    and generates multiple output formats
    """
    
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
    
    def plan_corrections(
        self,
        extracted_data: Dict,
        ground_truth: 'GroundTruth',
        validation_issues: List[str]
    ) -> Dict[str, any]:
        """
        Create an intelligent correction plan using Claude
        
        Returns a structured plan with:
        - Priority level for each correction
        - Specific values to change
        - Reasoning for each correction
        """
        
        prompt = f"""
        You are a QA specialist for HR job requisitions at Albertsons Companies.
        
        Analyze the following validation issues and create a detailed correction plan.
        
        Ground Truth (correct values):
        - Min Pay Rate: ${ground_truth.min_pay_rate:.2f}
        - Max Pay Rate: ${ground_truth.max_pay_rate:.2f}
        - Business Unit: {ground_truth.business_unit_name}
        - State: {ground_truth.state}
        
        Extracted Data (from document):
        {extracted_data}
        
        Validation Issues:
        {chr(10).join(f"{i+1}. {issue}" for i, issue in enumerate(validation_issues))}
        
        Create a correction plan with:
        1. Priority ranking (CRITICAL, HIGH, MEDIUM, LOW)
        2. Specific field changes needed
        3. Exact values to use
        4. Business justification for each correction
        
        Return as structured JSON:
        {{
            "corrections": [
                {{
                    "priority": "CRITICAL",
                    "field": "minimum_pay_rate",
                    "current_value": "...",
                    "corrected_value": "...",
                    "justification": "...",
                    "location_in_document": "..."
                }},
                ...
            ],
            "summary": "...",
            "risk_assessment": "..."
        }}
        """
        
        response = self.client.messages.create(
            model=self.model,
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )
        
        import json
        return json.loads(response.content[0].text)
    
    def generate_correction_report(
        self,
        job_id: str,
        correction_plan: Dict,
        applied_corrections: List[str]
    ) -> str:
        """
        Generate a human-readable correction report
        """
        
        report = f"""
╔══════════════════════════════════════════════════════════════╗
║        TA OPS AUTOMATED CORRECTION REPORT                     ║
╚══════════════════════════════════════════════════════════════╝

Job Requisition: {job_id}
Date: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Status: {"✓ CORRECTIONS APPLIED" if applied_corrections else "⚠ NO CORRECTIONS NEEDED"}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CORRECTION SUMMARY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{correction_plan.get('summary', 'No summary available')}

Risk Assessment: {correction_plan.get('risk_assessment', 'Low')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CORRECTIONS APPLIED
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        
        for i, correction in enumerate(correction_plan.get('corrections', []), 1):
            priority_symbol = {
                'CRITICAL': '🔴',
                'HIGH': '🟠',
                'MEDIUM': '🟡',
                'LOW': '🟢'
            }.get(correction['priority'], '⚪')
            
            report += f"""
{i}. {priority_symbol} {correction['priority']} - {correction['field']}
   
   Current Value:   {correction['current_value']}
   Corrected Value: {correction['corrected_value']}
   
   Location: {correction['location_in_document']}
   
   Justification:
   {correction['justification']}
   
   ───────────────────────────────────────────────────────────────
"""
        
        report += f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NEXT STEPS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. Review corrected document
2. Verify all changes against ground truth
3. Update Oracle HCM system if needed
4. Notify recruiting team of corrections
5. Archive original document for audit trail

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Generated by TA Ops Audit AI Agent
Philtech Inc. - GenAI Team
"""
        
        return report


# Integration example
if __name__ == "__main__":
    
    from ta_ops_audit_agent import GroundTruth
    
    # Example usage
    corrector = EnhancedDocumentCorrector(api_key="your-api-key")
    
    ground_truth = GroundTruth(
        job_requisition_number="651640",
        business_unit_name="27R-Seattle",
        min_pay_rate=14.75,
        max_pay_rate=15.00,
        facility="1148",
        state="IL"
    )
    
    # Job data (extracted from PDF)
    job_data = {
        'title': 'Retail Sales and Store Support',
        'job_id': '651640',
        'category': 'Retail, Store Ops',
        'posting_date': '11/04/2025, 11:58 PM',
        'schedule': 'Part time',
        'location': '1148 OGDEN AVE, DOWNERS GROVE, IL, 60515, US',
        'banner': 'Jewel Osco',
        'day_in_life': 'Our sales and store support teams...',
        'qualifications': [
            'You take pride in the work you do',
            'You agree that food is central to all our lives',
            # ... more items
        ],
        'benefits': [
            'Diverse and Inclusive work culture',
            'Competitive Wages paid weekly',
            # ... more items
        ],
        'role_rates': {
            'Grocery Associate': (13.75, 14.00),
            'Bakery Associate': (13.75, 14.00),
        }
    }
    
    pay_transparency = "Starting rates will be no less than the local minimum wage..."
    
    # Generate corrected document
    output = corrector.generate_corrected_document(
        job_data=job_data,
        ground_truth=ground_truth,
        pay_transparency_text=pay_transparency,
        output_path="/home/claude/651640_CORRECTED.docx"
    )
    
    print(f"✓ Corrected document generated: {output}")
